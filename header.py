from pathlib import Path
from docx import Document
import fitz  # PyMuPDF

# Step 1: Extract Text from PDF with font and layout info
def extract_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    all_text = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            lines = block.get("lines", [])
            for line in lines:
                spans = line.get("spans", [])
                for span in spans:
                    text = span["text"].strip()
                    if text:
                        all_text.append({
                            "text": text,
                            "size": span["size"],
                            "font": span["font"],
                            "flags": span["flags"]
                        })
    return all_text

# Step 2: Use heuristic to classify headings and paragraphs
def group_by_structure(elements):
    structured = []
    for e in elements:
        if e['size'] >= 14 and "Bold" in e['font']:
            structured.append({"type": "heading_1", "text": e['text']})
        elif e['size'] >= 13:
            structured.append({"type": "heading_2", "text": e['text']})
        else:
            structured.append({"type": "paragraph", "text": e['text']})
    return structured

# Step 3: Write to Word document
def write_to_word(structured_elements, output_path):
    doc = Document()
    for item in structured_elements:
        if item["type"] == "heading_1":
            doc.add_heading(item["text"], level=1)
        elif item["type"] == "heading_2":
            doc.add_heading(item["text"], level=2)
        else:
            doc.add_paragraph(item["text"])
    doc.save(output_path)

# Step 4: End-to-end function
def pdf_to_word(pdf_path, output_docx_path):
    elements = extract_pdf_text(pdf_path)
    structured = group_by_structure(elements)
    write_to_word(structured, output_docx_path)
    return output_docx_path

# Prepare sample output path
output_path = Path("/mnt/data/converted_output.docx")
output_path


pdf_to_word("your_file.pdf", "your_output.docx")
